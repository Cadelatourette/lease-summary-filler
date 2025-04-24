import streamlit as st
import fitz  # PyMuPDF
from docx import Document
from io import BytesIO
import re

def extract_text_from_pdf(file):
    text = ""
    with fitz.open(stream=file.read(), filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_lease_data(text):
    # Very basic sample logic – replace with more advanced extraction as needed
    data = {
        "{{Tenant_Name}}": re.search(r"Tenant Name[:\s]+(.+)", text, re.IGNORECASE),
        "{{Commencement_Date}}": re.search(r"Commencement Date[:\s]+(.+)", text, re.IGNORECASE),
        "{{Base_Rent}}": re.search(r"Base Rent[:\s]+(.+)", text, re.IGNORECASE),
        "{{Pro_Rata_Share}}": re.search(r"Pro Rata Share[:\s]+(.+)", text, re.IGNORECASE)
    }
    return {k: (v.group(1).strip() if v else "Not Provided") for k, v in data.items()}

def fill_template(template_file, data):
    doc = Document(template_file)
    for p in doc.paragraphs:
        for key, value in data.items():
            if key in p.text:
                p.text = p.text.replace(key, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

st.title("Lease Summary Filler")

lease_file = st.file_uploader("Upload Lease Document (PDF only for now)", type=["pdf"])
template_file = st.file_uploader("Upload Lease Summary Template (.docx)", type=["docx"])

if lease_file and template_file:
    with st.spinner("Extracting lease data..."):
        text = extract_text_from_pdf(lease_file)
        lease_data = extract_lease_data(text)

    with st.spinner("Filling template..."):
        result_doc = fill_template(template_file, lease_data)

    st.success("Lease summary ready!")
    st.download_button("Download Filled Lease Summary", result_doc, file_name="Lease_Summary_Filled.docx")
